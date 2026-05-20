import os
import re
import json

import pdfplumber
from groq import Groq
from dotenv import load_dotenv
from docx import Document

load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
client = Groq(api_key=GROQ_API_KEY)


# ---------------- CLEAN TEXT ---------------- #

def clean_text(text):
    text = text.replace('\u200b', ' ')
    text = text.replace('\u200c', ' ')
    text = text.replace('\u200d', ' ')
    text = text.replace('\ufeff', ' ')
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', ' ', text)
    text = re.sub(r' +', ' ', text)
    return text.strip()


# ---------------- PDF TEXT + TABLE EXTRACTION ---------------- #

def extract_text_from_pdf(file_path):
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # Extract normal text
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

                # Extract tables (Change 6)
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        row_text = " | ".join(
                            [cell if cell else "" for cell in row]
                        )
                        text += row_text + "\n"
    except Exception as e:
        text = ""
    return clean_text(text)


# ---------------- OCR FOR IMAGE-BASED PDF (Change 5) ---------------- #

def extract_text_from_image_pdf(file_path):
    """
    Uses OCR to extract text from image-based or scanned PDFs.
    Requires: pip install pdf2image pytesseract pillow
    Requires: Tesseract OCR installed on Windows
    Download: https://github.com/UB-Mannheim/tesseract/wiki
    """
    try:
        import pytesseract
        from pdf2image import convert_from_path

        # Set tesseract path for Windows
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

        images = convert_from_path(file_path, dpi=300)
        text = ""
        for image in images:
            page_text = pytesseract.image_to_string(image)
            text += page_text + "\n"

        return clean_text(text)

    except Exception as e:
        return ""


# ---------------- IMAGE FILE EXTRACTION (Change 5) ---------------- #

def extract_text_from_image(file_path):
    """
    Extracts text from direct image uploads (JPG, PNG etc.)
    """
    try:
        import pytesseract
        from PIL import Image

        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return clean_text(text)

    except Exception as e:
        return ""


# ---------------- DOCX TEXT EXTRACTION ---------------- #

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"

    # Extract tables from DOCX
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text for cell in row.cells])
            text += row_text + "\n"

    return clean_text(text)


# ---------------- SMART TEXT EXTRACTOR ---------------- #

def extract_text(file_path):
    """
    Smart extractor - handles normal PDF, image PDF (OCR), DOCX, images
    """
    ext = file_path.lower()

    if ext.endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
        # If very little text, it's image-based — use OCR
        if len(text.strip()) < 100:
            print(f"[INFO] Low text detected, switching to OCR: {file_path}")
            text = extract_text_from_image_pdf(file_path)
        return text

    elif ext.endswith(".docx"):
        return extract_text_from_docx(file_path)

    elif ext.endswith((".jpg", ".jpeg", ".png", ".webp", ".bmp")):
        return extract_text_from_image(file_path)

    else:
        return ""


# ---------------- GROQ EXTRACTION ---------------- #

def extract_resume_data(file_path):

    text = extract_text(file_path)

    if not text or len(text.strip()) < 50:
        return {
            "name": "NA", "email": "NA", "phone": "NA",
            "qualification": [], "passout_year": "NA",
            "key_skills": [], "years_of_experience": "NA",
            "previous_companies": [], "internships": [],
            "linkedin": "NA", "github": "NA", "projects": []
        }

    prompt = f"""
You are an advanced ATS Resume Parser for ALL resume types:
IT (Software Engineers, Data Scientists, Developers) and
Non-IT (Commerce, Banking, Finance, Marketing, HR, Arts, Science).

Extract details from the resume and return ONLY valid JSON.
No explanation, no markdown, no code fences.

Required fields:
{{
  "name": "",
  "email": "",
  "phone": "",
  "qualification": [],
  "passout_year": "",
  "key_skills": [],
  "years_of_experience": "",
  "previous_companies": [],
  "internships": [],
  "linkedin": "",
  "github": "",
  "projects": []
}}

Rules:

- If a field is missing return "NA"
- key_skills: array of strings. Include both IT skills (Python, React) and Non-IT skills (Accounting, GST, Tally, Banking, Marketing, Communication)
- projects: array of strings
- qualification: array of strings

- passout_year:
  * Completed degree → return graduation year as string e.g. "2023"
  * Currently studying / final year / last semester / end date is future year → return "Pursuing"
  * Not mentioned → return "NA"

- previous_companies:
  * Array of company names where person worked as EMPLOYEE (full-time or part-time)
  * Do NOT include internship companies here
  * Example: ["Infosys", "TCS"]
  * If none → return []

- internships:
  * Array of objects for each internship found
  * Format: [{{"company": "ABC Corp", "role": "Data Analyst Intern", "duration": "3 months"}}]
  * Include ALL internships even if short
  * If none → return []

- years_of_experience:
  * Total work experience e.g. "6 months", "2 years"
  * Students/freshers with no work experience → return "Fresher"
  * Calculate from date ranges if needed

- Return ONLY the JSON object, nothing else

Resume Text:
{text}
"""

    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=2000,
        )

        raw_output = response.choices[0].message.content.strip()
        raw_output = raw_output.replace("```json", "").replace("```", "").strip()

        data = json.loads(raw_output)
        return data

    except Exception as e:
        return {"error": str(e)}